"""Create bilingual DOCX fixtures for automatic equation classification evals."""

from __future__ import annotations

import argparse
import logging
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


LOGGER = logging.getLogger("auto-classification-fixtures")
BODY_SIZE_PT = 12
TITLE_SIZE_PT = 16
BODY_AFTER_PT = 8
LINE_SPACING = 1.25


ZH_PARAGRAPHS = (
    "量測雜訊變異數 [[MATH id=measurement_variance tex=\\sigma_v^2]] 用於描述重複量測的離散程度，其單位為 rad^2。",
    "扣除系統偏移量後的校正量可表示如下：",
    "[[MATH id=offset_correction tex=\\widetilde{y}_i = y_i - b]]",
    "其中，y_i 為第 i 個取樣位置之原始量測值，單位為 rad；b 為系統偏移量，單位為 rad；\\widetilde{y}_i 為校正後量測值，單位為 rad；i 為無因次取樣位置索引。",
    "單一取樣位置的偏折角增量可由下列模型求得：",
    "[[MATH id=local_deflection tex=\\Delta \\theta_{y,i} = K_i \\Delta n_i]]",
    "其中，\\Delta \\theta_{y,i} 為第 i 個取樣位置在 y 方向之偏折角增量，單位為 rad；K_i 為第 i 個位置之靈敏度係數，單位為 rad；\\Delta n_i 為無因次折射率差；i 為無因次取樣位置索引。",
    "由式 [[REF id=local_derivation target=local_deflection]] 得知，局部偏折角增量與折射率差成正比。",
    "為降低不同取樣數量造成的尺度差異，y 方向之累積偏折角定義如下：",
    "[[MATH id=cumulative_deflection tex=\\theta_y = \\frac{1}{N}\\sum_{i=1}^{N}\\Delta \\theta_{y,i}]]",
    "其中，\\theta_y 為 y 方向之累積偏折角，單位為 rad；\\Delta \\theta_{y,i} 為第 i 個取樣位置之偏折角增量，單位為 rad；i 為取樣位置索引；N 為取樣位置總數，i 與 N 皆為無因次正整數。",
    "其於 y 方向之累積偏折角可表示如式 [[REF id=cumulative_statement target=cumulative_deflection]] 所示。",
    "不同試次直接比較可能缺乏直觀的可比性。因此，如式 [[REF id=cumulative_comparison target=cumulative_deflection]] 所示，採用了依取樣位置總數正規化的累積偏折角。",
)

EN_PARAGRAPHS = (
    "The measurement-noise variance [[MATH id=measurement_variance tex=\\sigma_v^2]] describes repeatability and has units of rad^2.",
    "The offset-corrected measurement is expressed as follows:",
    "[[MATH id=offset_correction tex=\\widetilde{y}_i = y_i - b]]",
    "where y_i is the raw value at sample i in rad, b is the systematic offset in rad, \\widetilde{y}_i is the corrected value in rad, and i is a dimensionless sample index.",
    "The local deflection-angle increment is obtained from the following model:",
    "[[MATH id=local_deflection tex=\\Delta \\theta_{y,i} = K_i \\Delta n_i]]",
    "where \\Delta \\theta_{y,i} is the y-direction deflection-angle increment at sample i in rad, K_i is the sensitivity coefficient at sample i in rad, \\Delta n_i is the dimensionless refractive-index difference, and i is a dimensionless sample index.",
    "Equation [[REF id=local_derivation target=local_deflection]] indicates that the local deflection increment is proportional to the refractive-index difference.",
    "To reduce scale differences caused by unequal sample counts, the cumulative y-direction deflection angle is defined as follows:",
    "[[MATH id=cumulative_deflection tex=\\theta_y = \\frac{1}{N}\\sum_{i=1}^{N}\\Delta \\theta_{y,i}]]",
    "where \\theta_y is the cumulative y-direction deflection angle in rad, \\Delta \\theta_{y,i} is the deflection-angle increment at sample i in rad, i is the sample index, and N is the sample count; i and N are dimensionless positive integers.",
    "Its cumulative deflection angle in the y direction can be expressed as shown in Eq. [[REF id=cumulative_statement target=cumulative_deflection]].",
    "Direct comparison may lack an intuitive basis when sample counts differ. Therefore, as shown in Eq. [[REF id=cumulative_comparison target=cumulative_deflection]], the sample-count-normalized cumulative deflection angle was adopted.",
)


def set_run_font(run, western_font: str, east_asia_font: str, size_pt: int) -> None:
    run.font.name = western_font
    run.font.size = Pt(size_pt)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)


def create_fixture(
    output_path: Path,
    title: str,
    paragraphs: tuple[str, ...],
    western_font: str,
    east_asia_font: str,
) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = western_font
    normal.font.size = Pt(BODY_SIZE_PT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), western_font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), western_font)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER_PT)
    normal.paragraph_format.line_spacing = LINE_SPACING

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(16)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, western_font, east_asia_font, TITLE_SIZE_PT)
    title_run.bold = True

    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(BODY_AFTER_PT)
        paragraph.paragraph_format.line_spacing = LINE_SPACING
        run = paragraph.add_run(text)
        set_run_font(run, western_font, east_asia_font, BODY_SIZE_PT)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = output_path.with_suffix(".tmp.docx")
    try:
        document.save(temporary_path)
        os.replace(temporary_path, output_path)
    except Exception:
        LOGGER.exception("Failed to write fixture: %s", output_path)
        if temporary_path.exists():
            temporary_path.unlink()
        raise
    LOGGER.info("Created %s", output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    logging.basicConfig(level=logging.INFO, format="[fixture] %(message)s")
    args = parse_args()
    try:
        create_fixture(
            args.output_dir / "zh-auto-classification-draft.docx",
            "偏折角公式自動分類草稿",
            ZH_PARAGRAPHS,
            "Times New Roman",
            "新細明體",
        )
        create_fixture(
            args.output_dir / "en-auto-classification-draft.docx",
            "Automatic Equation Classification Draft",
            EN_PARAGRAPHS,
            "Times New Roman",
            "Times New Roman",
        )
    except Exception:
        LOGGER.exception("Fixture generation failed")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
